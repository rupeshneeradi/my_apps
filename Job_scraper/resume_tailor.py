"""
Auto resume tailor — v1.0.1

For every job that passes filters and gets scored:
  1. Loads the matched base resume DOCX from cache
  2. Clones it (original is never modified)
  3. Prepends a "TAILORED FOR THIS ROLE" block at the top:
       • Job title + company + ATS score + reference key
       • Ready-to-paste SENTENCES for each missing skill group
         (e.g. "Designed RESTful/SOAP web services using JSON/XML …")
       • A summary of your existing strengths (no changes needed there)
  4. Saves to  tailored_resumes/{job_id}.docx

HOW TO USE:
  1. Open tailored_resumes/{key}.docx
  2. Read the ⚡ "Add to your resume" sentences near the top
  3. Copy any that are true of your experience → paste into the relevant
     section of your resume (summary / skills / project bullets)
  4. Apply — the reference key in the email matches this filename

The base resume is NEVER changed — each output is a fresh clone.
"""
import io
import logging
import os

from docx import Document
from docx.shared import Pt, RGBColor

log = logging.getLogger(__name__)

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "tailored_resumes")


# ── Skill-sentence templates ──────────────────────────────────────────────────
# Each entry: (set_of_trigger_keywords, sentence_template)
# {terms} = comma-joined missing keywords from this group that appeared in the JD.
# Longest-first inside each set so greedy matching works correctly.

_SKILL_SENTENCES: list[tuple[set, str]] = [
    # ── OIC / Integration Cloud ───────────────────────────────────────────────
    ({"oracle oic", "oracle integration cloud", "oic", "oracle integration",
      "integration cloud", "integration cloud service", "oracle ipaas"},
     "Built end-to-end enterprise integrations using Oracle Integration Cloud ({terms}) "
     "with adapters, orchestration flows, and scheduled/real-time connectivity patterns."),

    # ── SOA / Middleware ──────────────────────────────────────────────────────
    ({"oracle soa", "soa suite", "bpel", "oracle bpel", "bpm",
      "oracle esb", "service bus", "oracle mft", "managed file transfer"},
     "Implemented service-oriented architecture solutions using {terms} "
     "for enterprise middleware and process automation."),

    # ── REST / SOAP / API ─────────────────────────────────────────────────────
    ({"rest api", "restful", "rest adapter", "soap api", "soap adapter",
      "web services", "wsdl", "api gateway", "oauth", "api integration",
      "api management", "api platform"},
     "Designed and consumed {terms} for secure, scalable system-to-system data exchange."),

    # ── JSON / XML ────────────────────────────────────────────────────────────
    ({"json", "xml", "xslt", "xsd"},
     "Processed and transformed {terms} payloads in integration and ETL pipelines."),

    # ── ORDS ──────────────────────────────────────────────────────────────────
    ({"ords", "oracle rest data services", "rest enabled sql"},
     "Published Oracle database objects as RESTful APIs via {terms}, "
     "enabling consumption by mobile and web clients."),

    # ── APEX features ─────────────────────────────────────────────────────────
    ({"oracle apex", "application express", "oracle application express",
      "apex developer", "apex consultant", "apex 24", "apex 23", "apex 22", "apex 21",
      "low code", "low-code", "rapid application development", "oracle low code"},
     "Developed enterprise-grade web applications on Oracle APEX ({terms}), "
     "delivering low-code solutions with complex business workflows."),

    ({"apex plugins", "apex dynamic actions", "apex components",
      "apex collections", "apex interactive report", "apex interactive grid",
      "apex themes", "universal theme", "jet charts", "apex security",
      "apex authentication", "apex authorization"},
     "Implemented advanced APEX UI components — {terms} — "
     "for intuitive, data-driven user experiences."),

    # ── JavaScript / Front-end ────────────────────────────────────────────────
    ({"javascript", "jquery", "ajax", "html5", "html", "css"},
     "Enhanced application interfaces using {terms} "
     "for dynamic client-side interactivity and responsive layouts."),

    # ── Python ────────────────────────────────────────────────────────────────
    ({"python", "python etl"},
     "Automated data processing and workflow orchestration using {terms} scripts."),

    # ── PL/SQL core ───────────────────────────────────────────────────────────
    ({"stored procedures", "stored procedure", "triggers", "packages",
      "functions", "pl/sql", "plsql"},
     "Authored complex {terms} encapsulating business rules, "
     "data validations, and batch processing logic."),

    # ── SQL tuning ────────────────────────────────────────────────────────────
    ({"performance tuning", "query optimization", "sql tuning",
      "execution plan", "explain plan", "index tuning",
      "partitioning", "indexing", "hints"},
     "Improved query performance through {terms}, "
     "reducing execution times by analysing and restructuring execution plans."),

    # ── Advanced PL/SQL ───────────────────────────────────────────────────────
    ({"bulk collect", "forall", "dynamic sql", "ref cursor",
      "cursor", "exception handling", "materialized views",
      "dbms_scheduler", "dbms_sql", "utl_file"},
     "Leveraged advanced PL/SQL constructs — {terms} — "
     "for high-volume data processing and runtime flexibility."),

    # ── Oracle EBS modules ────────────────────────────────────────────────────
    ({"oracle financials", "oracle ar", "oracle ap", "oracle gl",
      "oracle fa", "oracle om", "oracle inventory",
      "oracle purchasing", "oracle receivables", "oracle payables",
      "oracle general ledger", "oracle supply chain", "oracle procurement"},
     "Configured and customised Oracle EBS modules ({terms}) "
     "for end-to-end financial and supply-chain operations."),

    ({"oracle forms", "oracle reports", "oracle workflow",
      "oa framework", "oaf", "bi publisher", "xml publisher",
      "oracle concurrent program"},
     "Developed EBS technical components — {terms} — "
     "to support reporting, data entry, and automated processing."),

    # ── Oracle ERP Cloud / Fusion ─────────────────────────────────────────────
    ({"oracle erp cloud", "oracle cloud erp", "oracle cloud implementation",
      "oracle cloud consultant", "oracle cloud technical", "oracle saas",
      "fusion applications", "oracle cloud"},
     "Delivered {terms} implementations covering gap analysis, "
     "configuration, extensions, and go-live cutover."),

    ({"adf", "oracle adf", "adf bc", "adf faces", "jdeveloper",
      "fusion web client"},
     "Built Oracle Fusion UI extensions and custom pages using {terms}."),

    ({"oracle vbcs", "visual builder"},
     "Created lightweight web apps and extensions using {terms} "
     "on Oracle Cloud with no-server deployment."),

    # ── BI / Reporting ────────────────────────────────────────────────────────
    ({"otbi", "oracle transactional bi", "oracle bi publisher", "bi publisher",
      "faw", "oracle analytics cloud", "oracle analytics",
      "financial reporting studio"},
     "Delivered management and operational reports using {terms} "
     "for real-time financial and operational visibility."),

    # ── HCM-specific ──────────────────────────────────────────────────────────
    ({"fast formula", "fastformula"},
     "Wrote {terms} rules for payroll calculations, "
     "absence policies, and eligibility criteria."),

    ({"hcm extracts", "hcm data loader", "hdl",
      "hcm spreadsheet loader", "hdi"},
     "Managed bulk data using {terms} for mass-load, "
     "migration, and reconciliation across HCM modules."),

    ({"flexfields", "sandbox"},
     "Extended Oracle Cloud HCM data model using {terms} "
     "for client-specific configurations without custom code."),

    # ── NetSuite scripting ────────────────────────────────────────────────────
    ({"suitescript 2.0", "suitescript 1.0", "suitescript",
      "suitelet", "restlet", "user event script",
      "scheduled script", "client script", "map reduce",
      "workflow action script"},
     "Developed NetSuite business logic using {terms} "
     "for custom workflows, UI interactions, and background processing."),

    ({"netsuite workflow", "netsuite customization",
      "netsuite saved search", "custom fields", "custom records"},
     "Delivered end-user NetSuite customisations — {terms} — "
     "without bespoke code where configuration sufficed."),

    ({"netsuite integration", "netsuite connector", "celigo", "boomi netsuite"},
     "Integrated NetSuite with third-party platforms using {terms} "
     "for bi-directional data synchronisation."),

    # ── Informatica ───────────────────────────────────────────────────────────
    ({"informatica powercenter", "informatica iics", "informatica bdm",
      "informatica cloud", "informatica idmc", "informatica", "powercenter"},
     "Designed and optimised ETL workflows using {terms} "
     "for enterprise data integration and migration projects."),

    # ── DataStage ─────────────────────────────────────────────────────────────
    ({"ibm datastage", "datastage parallel", "datastage server", "datastage"},
     "Built high-throughput parallel ETL jobs using {terms} "
     "for large-scale data warehouse loading."),

    # ── ODI ───────────────────────────────────────────────────────────────────
    ({"oracle data integrator", "odi"},
     "Developed data integration mappings and interfaces using {terms} "
     "for Oracle-centric data warehouse environments."),

    # ── Other ETL tools ───────────────────────────────────────────────────────
    ({"talend", "talend open studio", "talend cloud",
      "pentaho", "kettle",
      "ssis", "sql server integration services"},
     "Created data pipelines using {terms} for cross-platform ETL delivery."),

    # ── Cloud ETL ─────────────────────────────────────────────────────────────
    ({"azure data factory", "adf pipeline", "azure synapse"},
     "Implemented cloud data integration pipelines using {terms} "
     "on Microsoft Azure."),

    ({"aws glue", "aws etl", "aws data pipeline"},
     "Built serverless ETL solutions using {terms} on AWS."),

    ({"apache spark", "pyspark", "databricks"},
     "Processed large-scale datasets using {terms} "
     "for distributed in-memory transformations."),

    ({"apache kafka", "kafka"},
     "Streamed real-time events using {terms} for high-throughput data pipelines."),

    ({"snowflake", "redshift", "bigquery", "big query"},
     "Loaded and queried cloud data warehouse platforms — {terms} — "
     "for analytical workloads."),

    ({"dbt", "data build tool"},
     "Transformed and documented data models in the warehouse using {terms}."),

    ({"apache airflow", "airflow", "nifi", "apache nifi"},
     "Orchestrated data pipeline schedules and dependencies using {terms}."),

    # ── Data modelling ────────────────────────────────────────────────────────
    ({"dimensional modeling", "star schema", "snowflake schema",
      "data vault", "slowly changing dimension", "scd",
      "fact table", "dimension table", "data modeling"},
     "Designed warehouse schemas using {terms} "
     "to optimise query performance and historical tracking."),

    # ── Shell / scripting ─────────────────────────────────────────────────────
    ({"shell scripting", "unix scripting"},
     "Automated operational tasks and job scheduling using {terms}."),

    # ── DevOps / process ──────────────────────────────────────────────────────
    ({"agile", "scrum"},
     "Delivered work in {terms} sprints, participating in stand-ups, "
     "retrospectives, and sprint planning."),

    ({"git", "jenkins", "devops"},
     "Managed source code and deployments using {terms} "
     "for version control and continuous integration."),

    ({"jira"},
     "Tracked tasks, bugs, and sprint progress using {terms}."),

    # ── Oracle infra ──────────────────────────────────────────────────────────
    ({"oracle rac", "oracle exadata", "data pump", "oracle goldengate", "goldengate"},
     "Administered and tuned high-availability Oracle infrastructure — {terms}."),
]


# ── helpers ───────────────────────────────────────────────────────────────────

def _wb_match(keyword: str, text_lower: str) -> bool:
    """Simple word-boundary check without regex import."""
    import re
    return bool(re.search(r"\b" + re.escape(keyword) + r"\b", text_lower))


# ── Display-name overrides for keyword formatting ─────────────────────────────
_DISPLAY: dict[str, str] = {
    "rest api":               "REST API",
    "restful":                "RESTful",
    "soap api":               "SOAP API",
    "wsdl":                   "WSDL",
    "xslt":                   "XSLT",
    "xsd":                    "XSD",
    "json":                   "JSON",
    "xml":                    "XML",
    "html5":                  "HTML5",
    "html":                   "HTML",
    "css":                    "CSS",
    "javascript":             "JavaScript",
    "jquery":                 "jQuery",
    "ajax":                   "AJAX",
    "pl/sql":                 "PL/SQL",
    "plsql":                  "PL/SQL",
    "sql":                    "SQL",
    "ssis":                   "SSIS",
    "scd":                    "SCD",
    "etl":                    "ETL",
    "oic":                    "OIC",
    "ics":                    "ICS",
    "ords":                   "ORDS",
    "adf":                    "ADF",
    "bpel":                   "BPEL",
    "bpm":                    "BPM",
    "git":                    "Git",
    "jira":                   "Jira",
    "aws glue":               "AWS Glue",
    "aws etl":                "AWS ETL",
    "aws data pipeline":      "AWS Data Pipeline",
    "dbt":                    "dbt",
    "pyspark":                "PySpark",
    "bigquery":               "BigQuery",
    "big query":              "BigQuery",
    "nifi":                   "Apache NiFi",
    "apache nifi":            "Apache NiFi",
    "apache kafka":           "Apache Kafka",
    "apache airflow":         "Apache Airflow",
    "apache spark":           "Apache Spark",
    "hdl":                    "HDL (HCM Data Loader)",
    "hdi":                    "HDI",
    "otl":                    "OTL",
    "otbi":                   "OTBI",
    "faw":                    "FAW",
    "odi":                    "ODI",
    "vbcs":                   "VBCS",
    "oracle vbcs":            "Oracle VBCS",
    "oracle adf":             "Oracle ADF",
    "oracle soa":             "Oracle SOA",
    "oaf":                    "OA Framework",
    "bi publisher":           "BI Publisher",
    "oracle bi publisher":    "Oracle BI Publisher",
    "suitescript 2.0":        "SuiteScript 2.0",
    "suitescript 1.0":        "SuiteScript 1.0",
    "suitescript":            "SuiteScript",
    "suitelet":               "Suitelet",
    "restlet":                "RESTlet",
    "datastage":              "DataStage",
    "ibm datastage":          "IBM DataStage",
    "datastage parallel":     "DataStage Parallel",
    "datastage server":       "DataStage Server",
    "netsuite":               "NetSuite",
    "oracle netsuite":        "Oracle NetSuite",
    "netsuite integration":   "NetSuite Integration",
    "netsuite connector":     "NetSuite Connector",
    "netsuite workflow":      "NetSuite Workflow",
    "netsuite customization": "NetSuite Customization",
    "netsuite saved search":  "NetSuite Saved Search",
    "powercenter":            "PowerCenter",
    "informatica powercenter":"Informatica PowerCenter",
    "informatica iics":       "Informatica IICS",
    "informatica bdm":        "Informatica BDM",
    "informatica cloud":      "Informatica Cloud",
    "informatica idmc":       "Informatica IDMC",
    "goldengate":             "GoldenGate",
    "oracle goldengate":      "Oracle GoldenGate",
    "oracle rac":             "Oracle RAC",
    "oracle exadata":         "Oracle Exadata",
    "jdeveloper":             "JDeveloper",
    "webcenter":              "WebCenter",
    "snowflake":              "Snowflake",
    "databricks":             "Databricks",
    "redshift":               "Redshift",
    "kafka":                  "Kafka",
    "airflow":                "Airflow",
    "celigo":                 "Celigo",
    "talend":                 "Talend",
    "pentaho":                "Pentaho",
    "kettle":                 "Kettle",
}


def _display_kw(kw: str) -> str:
    """Return a nicely formatted version of a keyword for use in sentences."""
    lo = kw.lower()
    if lo in _DISPLAY:
        return _DISPLAY[lo]
    # All-caps for short abbreviations (≤4 chars, no spaces)
    if len(kw) <= 4 and " " not in kw:
        return kw.upper()
    return kw.title()


def _build_missing_sentences(missing: list[str]) -> list[str]:
    """
    Convert a list of raw missing JD keywords into ready-to-paste resume sentences.
    Groups keywords by skill category; each group produces one sentence.
    Returns list of sentence strings.
    """
    missing_lower = {kw.lower() for kw in missing}
    used: set[str] = set()
    sentences: list[str] = []

    for trigger_set, template in _SKILL_SENTENCES:
        matched_in_group = []
        for kw in trigger_set:
            if kw in missing_lower and kw not in used:
                matched_in_group.append(kw)
        if not matched_in_group:
            continue
        # Mark as used so they don't appear in multiple sentences
        used.update(matched_in_group)
        # Format terms with proper display names
        terms_str = ", ".join(_display_kw(t) for t in matched_in_group)
        sentences.append(template.replace("{terms}", terms_str))

    # Any leftover keywords not caught by a template → plain list sentence
    leftover = [kw for kw in missing if kw.lower() not in used]
    if leftover:
        lo_str = ", ".join(leftover[:10])
        suffix = f" (+{len(leftover)-10} more)" if len(leftover) > 10 else ""
        sentences.append(
            f"Additionally, consider weaving in: {lo_str}{suffix}."
        )

    return sentences


def _group_matched(matched: list[str]) -> str:
    """Summarise matched keywords into a readable strength line."""
    if not matched:
        return "—"
    shown = matched[:20]
    rest  = len(matched) - 20
    out   = ", ".join(shown)
    if rest > 0:
        out += f"  (+{rest} more)"
    return out


# ── DOCX paragraph helpers ───────────────────────────────────────────────────

def _para(doc: Document, text: str, *,
          bold: bool = False,
          italic: bool = False,
          size: int = 10,
          rgb: tuple = (0, 0, 0)) -> None:
    """
    Create a paragraph and insert it at position 0 in the body.
    Builds bottom-up — last call ends up at the top of the document.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.bold        = bold
    run.italic      = italic
    run.font.size   = Pt(size)
    run.font.color.rgb = RGBColor(*rgb)

    body = doc.element.body
    body.remove(p._element)
    body.insert(0, p._element)


def _divider(doc: Document) -> None:
    _para(doc, "─" * 80, size=7, rgb=(190, 190, 190))


# ── Main header builder ───────────────────────────────────────────────────────

def _prepend_header(doc: Document, job: dict,
                    matched: list[str], missing: list[str]) -> None:
    """
    Insert the tailored coaching block BEFORE all existing resume content.
    Built bottom-up (each _para() inserts at position 0), so the last
    _para() call lands at the very top of the document.
    """
    ats     = job.get("ats_result")
    score   = f"{ats.score:.0f}%" if ats else "—"
    label   = ats.label           if ats else "—"
    title   = job.get("title",   "")
    company = job.get("company", "")
    portal  = job.get("portal",  "")
    key     = job.get("id",      "")
    url     = job.get("url",     "")

    missing_sentences = _build_missing_sentences(missing)

    # ── Build REVERSE order — last line written = first line in document ──────

    # Blank spacer after the block
    _para(doc, "", size=6)

    # Bottom divider
    _divider(doc)

    # ── Strengths section ──
    _para(doc,
          f"  ✓  {_group_matched(matched)}",
          bold=False, size=9, rgb=(20, 110, 20))
    _para(doc,
          "YOUR RESUME ALREADY COVERS:",
          bold=True, size=9, rgb=(20, 110, 20))

    # Spacer between sections
    _para(doc, "", size=5)

    # ── Missing-skill sentences section ──
    if missing_sentences:
        # Print in reverse so they appear top→bottom in doc
        for sentence in reversed(missing_sentences):
            _para(doc,
                  f"  •  {sentence}",
                  bold=False, italic=False, size=9, rgb=(160, 40, 0))
        _para(doc,
              "⚡ ADD TO YOUR RESUME  (copy, edit to match your experience, paste):",
              bold=True, size=9, rgb=(160, 40, 0))
    else:
        _para(doc,
              "  ✓  No significant gaps detected — your resume covers this JD well.",
              bold=True, size=9, rgb=(20, 110, 20))

    # Spacer
    _para(doc, "", size=5)

    # ── Instructions line ──
    _para(doc,
          "INSTRUCTIONS:  Review bullet points below → paste relevant ones into your "
          "Summary / Skills / Project sections → apply using the link above.",
          bold=False, italic=True, size=8, rgb=(80, 80, 80))

    # Top divider
    _divider(doc)

    # ── Meta / reference line ──
    _para(doc,
          f"Key: {key}   |   Source: {portal}   |   ATS Score: {score} [{label}]   |   {url}",
          bold=False, size=8, rgb=(110, 110, 110))

    # ── Heading ──
    _para(doc,
          f"TAILORED FOR: {title}  @  {company}",
          bold=True, size=12, rgb=(0, 65, 130))

    # Top-most divider
    _divider(doc)


# ── Public API ────────────────────────────────────────────────────────────────

def generate(job: dict, docx_bytes: bytes) -> str | None:
    """
    Clone the base resume, inject the tailored header, save.
    Returns the output path, or None on failure.
    """
    job_id = job.get("id", "unknown")
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    out_path = os.path.join(OUTPUT_DIR, f"{job_id}.docx")

    ats     = job.get("ats_result")
    matched = list(ats.matched) if ats else []
    missing = list(ats.missing) if ats else []

    try:
        doc = Document(io.BytesIO(docx_bytes))
        _prepend_header(doc, job, matched, missing)
        doc.save(out_path)
        log.debug("Tailored resume → %s", out_path)
        return out_path
    except Exception as exc:
        log.error("resume_tailor: failed for job %s (%s @ %s): %s",
                  job_id, job.get("title"), job.get("company"), exc)
        return None


def tailor_all(jobs: list[dict], resume_docx: dict[str, bytes]) -> int:
    """
    Generate tailored resumes for every qualifying job.

    Sets  job["tailored_resume"] = "{job_id}.docx"  on success.
    Returns number of files written.
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    count = 0

    for job in jobs:
        ats = job.get("ats_result")
        if not ats:
            continue

        base = ats.best_resume
        if base not in resume_docx:
            log.warning("No cached DOCX for '%s' — skipping tailor for %s",
                        base, job.get("id"))
            continue

        path = generate(job, resume_docx[base])
        if path:
            job["tailored_resume"] = os.path.basename(path)
            count += 1

    log.info("resume_tailor: %d tailored resumes written to %s", count, OUTPUT_DIR)
    return count
