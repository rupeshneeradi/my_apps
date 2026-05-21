"""
Resume Builder — generates a clean, ATS-friendly .docx resume from structured data.
No AI, no external calls, 100% local.
Single-column layout — no tables, no text boxes — passes every ATS parser.
"""
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _rgb(h: str) -> RGBColor:
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


BLACK  = _rgb("111827")
DARK   = _rgb("1f2937")
GRAY   = _rgb("4b5563")
LGRAY  = _rgb("9ca3af")
ACCENT = _rgb("1d4ed8")   # deep blue — professional, ATS-safe


def _hr(doc):
    """Thin coloured line under section headings."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "bfdbfe")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)


def _section(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(0)
    _hr(doc)


def _exp_head(doc, title: str, company: str, start: str, end: str, location: str = ""):
    p = doc.add_paragraph()
    r_t = p.add_run(title)
    r_t.bold = True
    r_t.font.size = Pt(11)
    r_t.font.color.rgb = DARK
    if company:
        r_s = p.add_run(f"  |  {company}")
        r_s.font.size = Pt(11)
        r_s.font.color.rgb = GRAY
    date_str = f"{start} – {end}" if start and end else (start or end or "")
    if date_str:
        r_d = p.add_run(f"  |  {date_str}")
        r_d.font.size = Pt(10)
        r_d.font.color.rgb = LGRAY
    if location:
        r_l = p.add_run(f"  •  {location}")
        r_l.font.size = Pt(10)
        r_l.font.color.rgb = LGRAY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)


def _bullet(doc, text: str):
    text = text.strip().lstrip("•–- ")
    if not text:
        return
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.2)


def generate_resume_docx(data: dict) -> bytes:
    """
    Expected data shape:
        personal      : {name, email, phone, location, linkedin, github, website}
        summary       : str
        experience    : [{title, company, start, end, location, bullets:[str]}]
        education     : [{degree, school, year, gpa}]
        skills        : [{category, items:[str]}]
        certifications: [{name, issuer, year}]
        projects      : [{name, tech, description}]   (optional)
    """
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.left_margin   = Inches(0.75)
    sec.right_margin  = Inches(0.75)
    sec.top_margin    = Inches(0.65)
    sec.bottom_margin = Inches(0.65)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    personal = data.get("personal") or {}

    # ── Name ──────────────────────────────────────────────────────────────────
    name = (personal.get("name") or "").strip()
    if name:
        p = doc.add_paragraph()
        r = p.add_run(name)
        r.bold = True
        r.font.size = Pt(20)
        r.font.color.rgb = DARK
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)

    # ── Contact line ──────────────────────────────────────────────────────────
    parts = []
    for f in ["email", "phone", "location"]:
        v = (personal.get(f) or "").strip()
        if v:
            parts.append(v)
    for f in ["linkedin", "github", "website"]:
        v = (personal.get(f) or "").strip()
        if v:
            parts.append(v)
    if parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("  •  ".join(parts))
        r.font.size = Pt(9.5)
        r.font.color.rgb = GRAY
        p.paragraph_format.space_after = Pt(4)

    # ── Summary ───────────────────────────────────────────────────────────────
    summary = (data.get("summary") or "").strip()
    if summary:
        _section(doc, "Professional Summary")
        p = doc.add_paragraph()
        p.add_run(summary).font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(3)

    # ── Experience ────────────────────────────────────────────────────────────
    experience = data.get("experience") or []
    if experience:
        _section(doc, "Professional Experience")
        for exp in experience:
            if not (exp.get("title") or exp.get("company")):
                continue
            _exp_head(doc,
                      exp.get("title", ""),
                      exp.get("company", ""),
                      exp.get("start", ""),
                      exp.get("end", "Present"),
                      exp.get("location", ""))
            for b in (exp.get("bullets") or []):
                _bullet(doc, b)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Education ─────────────────────────────────────────────────────────────
    education = data.get("education") or []
    if education:
        _section(doc, "Education")
        for edu in education:
            if not edu.get("degree"):
                continue
            p = doc.add_paragraph()
            r = p.add_run(edu.get("degree", ""))
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = DARK
            school = edu.get("school", "")
            if school:
                p.add_run(f"  —  {school}").font.size = Pt(11)
            year = edu.get("year", "")
            if year:
                p.add_run(f"  ({year})").font.size = Pt(10)
            gpa = edu.get("gpa", "")
            if gpa:
                p.add_run(f"  GPA: {gpa}").font.size = Pt(10)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(2)

    # ── Skills ────────────────────────────────────────────────────────────────
    skills = data.get("skills") or []
    if skills:
        _section(doc, "Technical Skills")
        for cat in skills:
            if not cat:
                continue
            if isinstance(cat, str):
                p = doc.add_paragraph()
                p.add_run(cat).font.size = Pt(10.5)
                continue
            cat_name = (cat.get("category") or "").strip()
            items    = [str(i).strip() for i in (cat.get("items") or []) if str(i).strip()]
            if not items:
                continue
            p = doc.add_paragraph()
            if cat_name:
                r = p.add_run(f"{cat_name}: ")
                r.bold = True
                r.font.size = Pt(10.5)
                r.font.color.rgb = ACCENT
            r2 = p.add_run(",  ".join(items))
            r2.font.size = Pt(10.5)
            r2.font.color.rgb = DARK
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)

    # ── Certifications ────────────────────────────────────────────────────────
    certifications = data.get("certifications") or []
    if certifications:
        _section(doc, "Certifications")
        for cert in certifications:
            if not cert.get("name"):
                continue
            p = doc.add_paragraph()
            r = p.add_run(cert.get("name", ""))
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = DARK
            issuer = cert.get("issuer", "")
            if issuer:
                p.add_run(f"  —  {issuer}").font.size = Pt(10.5)
            year = cert.get("year", "")
            if year:
                p.add_run(f"  ({year})").font.size = Pt(10)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(1)

    # ── Projects (optional) ───────────────────────────────────────────────────
    projects = data.get("projects") or []
    if projects:
        _section(doc, "Projects")
        for proj in projects:
            if not proj.get("name"):
                continue
            p = doc.add_paragraph()
            r = p.add_run(proj.get("name", ""))
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = DARK
            tech = (proj.get("tech") or "").strip()
            if tech:
                p.add_run(f"  |  {tech}").font.size = Pt(10)
            desc = (proj.get("description") or "").strip()
            if desc:
                _bullet(doc, desc)
            p.paragraph_format.space_after = Pt(2)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
