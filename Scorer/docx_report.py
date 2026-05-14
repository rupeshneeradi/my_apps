"""
docx_report.py — Generate a professional Word document resume improvement report.

Produces a fully formatted .docx with:
  • Score summary + verdict
  • Strengths and red flags
  • Ready-to-paste summary, skills, and experience bullets
  • Real HireITPeople examples per keyword
  • Keyword gap analysis
"""

import io
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colour palette ─────────────────────────────────────────────────────────────

def _rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

PURPLE = _rgb("7c3aed")
GREEN  = _rgb("15803d")
RED    = _rgb("dc2626")
ORANGE = _rgb("d97706")
BLUE   = _rgb("1d4ed8")
TEAL   = _rgb("0d9488")
GRAY   = _rgb("6b7280")
LGRAY  = _rgb("9ca3af")
BLACK  = _rgb("111827")

VERDICT_COLORS = {
    "STRONG FIT": GREEN,
    "GOOD FIT":   BLUE,
    "POSSIBLE":   ORANGE,
    "NOT A FIT":  RED,
}


# ── Low-level helpers ──────────────────────────────────────────────────────────

def _shade_paragraph(paragraph, hex_color: str):
    """Add a background fill shading to a paragraph via XML."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    pPr.append(shd)


def _section_heading(doc, text: str, color: RGBColor = PURPLE):
    p = doc.add_paragraph()
    _shade_paragraph(p, "f5f3ff")
    run = p.add_run(f"  {text}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def _sub_heading(doc, text: str, color: RGBColor = BLUE):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p


def _body(doc, text: str, color: RGBColor = BLACK,
          bold: bool = False, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.bold   = bold
    run.italic = italic
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)


def _bullet_para(doc, text: str, color: RGBColor = BLACK, size: int = 11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)


def _divider(doc):
    p = doc.add_paragraph()
    run = p.add_run("─" * 80)
    run.font.color.rgb = _rgb("e5e7eb")
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)


def _score_block(doc, ats: int, rec: int, qual: int, chance: str, verdict: str):
    """Inline score summary as a single formatted paragraph."""
    # Verdict line
    p = doc.add_paragraph()
    p.add_run("Verdict: ").bold = True
    vrun = p.add_run(f"  {verdict}  ")
    vrun.bold = True
    vrun.font.size = Pt(13)
    vrun.font.color.rgb = VERDICT_COLORS.get(verdict, BLACK)

    # Scores line
    p2 = doc.add_paragraph()
    for label, val, color in [
        ("ATS Match", f"{ats}%",      BLUE),
        ("Recruiter", f"{rec}/100",   PURPLE),
        ("Quality",   f"{qual}/100",  TEAL),
        ("Interview", chance,         ORANGE),
    ]:
        run = p2.add_run(f"  {label}: ")
        run.font.size = Pt(10)
        run.font.color.rgb = GRAY
        vr = p2.add_run(f"{val}   ")
        vr.bold = True
        vr.font.size = Pt(10)
        vr.font.color.rgb = color


# ── Public API ─────────────────────────────────────────────────────────────────

def generate_report_docx(
    resume_name: str,
    jd_title: str,
    analysis: dict,
    resume_content: dict,
    real_examples: dict,
) -> bytes:
    """
    Build and return a Word document (.docx bytes) containing the full
    resume improvement report for one resume vs one JD.
    """
    doc = Document()

    # ── Margins ───────────────────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.left_margin   = Inches(1.0)
    sec.right_margin  = Inches(1.0)
    sec.top_margin    = Inches(0.75)
    sec.bottom_margin = Inches(0.75)

    # ── Title block ───────────────────────────────────────────────────────────
    title_p = doc.add_heading("Resume Improvement Report", 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_p.runs:
        run.font.color.rgb = PURPLE
        run.font.size = Pt(22)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = sub_p.add_run(f"{resume_name}")
    r1.bold = True
    r1.font.size = Pt(13)
    r1.font.color.rgb = BLACK

    if jd_title:
        pos_p = doc.add_paragraph()
        pos_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rp = pos_p.add_run(f"Position: {jd_title}")
        rp.font.size = Pt(11)
        rp.font.color.rgb = GRAY

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = date_p.add_run(f"Generated {date.today().strftime('%B %d, %Y')}")
    rd.font.size = Pt(9)
    rd.font.color.rgb = LGRAY

    _divider(doc)

    # ── Score summary ─────────────────────────────────────────────────────────
    _section_heading(doc, "📊  Score Summary")
    _score_block(
        doc,
        analysis.get("ats_score", 0),
        analysis.get("recruiter_score", 0),
        analysis.get("quality_score", 0),
        analysis.get("interview_chance", "N/A"),
        analysis.get("verdict", ""),
    )

    # Strengths
    strengths = analysis.get("strengths", [])
    if strengths:
        doc.add_paragraph()
        _sub_heading(doc, "✓  Strengths", GREEN)
        for s in strengths:
            _bullet_para(doc, s, GREEN)

    # Red flags
    red_flags = analysis.get("red_flags", [])
    if red_flags:
        _sub_heading(doc, "⚠  Red Flags", RED)
        for f in red_flags:
            _bullet_para(doc, f, RED)

    _divider(doc)

    # ── Resume Writer ─────────────────────────────────────────────────────────
    _section_heading(doc, "✍  Ready-to-Use Resume Content", PURPLE)
    _body(doc, "Copy and paste these sections directly into your resume.",
          GRAY, italic=True, size=10)
    doc.add_paragraph()

    rc = resume_content or {}

    # 1 — Summary
    if rc.get("summary"):
        _sub_heading(doc, "1.  Professional Summary — Replace Your Current Summary", BLUE)
        _body(doc, "Add this at the very top of your resume:", GRAY, italic=True, size=10)
        p = doc.add_paragraph()
        _shade_paragraph(p, "eff6ff")
        run = p.add_run("  " + rc["summary"])
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK
        doc.add_paragraph()

    # 2 — Skills
    if rc.get("skills_to_add"):
        _sub_heading(doc, "2.  Technical Skills — Add These Lines", BLUE)
        _body(doc,
              "Add verbatim to your Technical Skills section "
              "(ATS matches strings literally):",
              GRAY, italic=True, size=10)
        for skill in rc["skills_to_add"]:
            _bullet_para(doc, skill, BLACK)
        doc.add_paragraph()

    # 3 — Exp 1
    if rc.get("exp1_bullets"):
        num_exp = rc.get("exp_blocks_detected", 1)
        label = (
            "3.  Most Recent Role — Add These Bullet Points"
            if num_exp >= 2 else
            "3.  Experience Section — Add These Bullet Points"
        )
        _sub_heading(doc, label, GREEN)
        _body(doc,
              "Add 1–2 of these below your existing bullets in your current role:",
              GRAY, italic=True, size=10)
        doc.add_paragraph()

        for b in rc["exp1_bullets"]:
            kw = b.get("keyword", "")
            _body(doc, f"▶  Keyword: {kw.upper()}", PURPLE, bold=True, size=10)

            if b.get("recruiter_why"):
                _body(doc, f"💡  {b['recruiter_why']}", GRAY, italic=True, size=10)

            p = doc.add_paragraph()
            _shade_paragraph(p, "f0fdf4")
            run = p.add_run("  • " + b.get("bullet", ""))
            run.font.size = Pt(11)
            run.font.color.rgb = BLACK

            # Real examples
            kw_key = kw.lower()
            examples = real_examples.get(kw_key) or real_examples.get(kw) or []
            if examples:
                _body(doc, "  📌 Real examples from hireitpeople.com:",
                      ORANGE, bold=True, size=9)
                for ex in examples:
                    p_ex = doc.add_paragraph()
                    run_ex = p_ex.add_run("      • " + ex)
                    run_ex.font.size = Pt(10)
                    run_ex.font.color.rgb = GRAY
                    run_ex.italic = True

            if b.get("alt_bullet"):
                _body(doc, "  Alternative version:", LGRAY, italic=True, size=9)
                p_alt = doc.add_paragraph()
                run_alt = p_alt.add_run("  • " + b["alt_bullet"])
                run_alt.font.size = Pt(10)
                run_alt.font.color.rgb = LGRAY

            doc.add_paragraph()

    # 4 — Exp 2
    if rc.get("exp2_bullets"):
        _sub_heading(doc, "4.  Previous Role — Add These Bullet Points", GREEN)
        _body(doc, "Add 1 of these to your previous role:", GRAY, italic=True, size=10)
        doc.add_paragraph()

        for b in rc["exp2_bullets"]:
            kw = b.get("keyword", "")
            _body(doc, f"▶  Keyword: {kw.upper()}", PURPLE, bold=True, size=10)

            if b.get("recruiter_why"):
                _body(doc, f"💡  {b['recruiter_why']}", GRAY, italic=True, size=10)

            p = doc.add_paragraph()
            _shade_paragraph(p, "f0fdf4")
            run = p.add_run("  • " + b.get("bullet", ""))
            run.font.size = Pt(11)
            run.font.color.rgb = BLACK

            kw_key = kw.lower()
            examples = real_examples.get(kw_key) or real_examples.get(kw) or []
            if examples:
                _body(doc, "  📌 Real examples from hireitpeople.com:",
                      ORANGE, bold=True, size=9)
                for ex in examples:
                    p_ex = doc.add_paragraph()
                    run_ex = p_ex.add_run("      • " + ex)
                    run_ex.font.size = Pt(10)
                    run_ex.font.color.rgb = GRAY
                    run_ex.italic = True

            doc.add_paragraph()

    # 5 — Keyword analysis
    _divider(doc)
    _section_heading(doc, "📋  Keyword Gap Analysis", PURPLE)

    gb = analysis.get("gap_breakdown", {})
    matched = analysis.get("matched", [])

    if gb.get("critical"):
        _body(doc, "🔴  Critical Gaps (appear in job title):", RED, bold=True, size=10)
        _body(doc, "  " + ", ".join(gb["critical"]), RED, size=10)

    if gb.get("required"):
        _body(doc, "🟡  Required (marked as required in JD):", ORANGE, bold=True, size=10)
        _body(doc, "  " + ", ".join(gb["required"]), ORANGE, size=10)

    if gb.get("preferred"):
        _body(doc, "🟢  Preferred (nice to have):", GRAY, bold=True, size=10)
        _body(doc, "  " + ", ".join(gb["preferred"]), GRAY, size=10)

    if matched:
        doc.add_paragraph()
        _body(doc, "✓  Keywords Already Present:", GREEN, bold=True, size=10)
        _body(doc, "  " + ", ".join(matched), GREEN, size=10)

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    _divider(doc)
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        "Generated by ATS Resume Scorer  •  Real examples sourced from hireitpeople.com"
    )
    fr.font.color.rgb = LGRAY
    fr.font.size = Pt(8)

    # ── Serialise ─────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
